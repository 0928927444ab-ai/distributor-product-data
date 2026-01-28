import pandas as pd
from docx import Document

# 讀取 Excel 檔案
print("=" * 60)
print("【樂莫集團稅務制度】樂莫集團.xlsx")
print("=" * 60)
try:
    df = pd.read_excel(r'C:\Users\User\Desktop\樂莫集團稅務制度\樂莫集團.xlsx', sheet_name=None)
    for name, sheet in df.items():
        print(f"\n--- Sheet: {name} ---")
        print(sheet.to_string())
except Exception as e:
    print(f"Error: {e}")

print("\n" + "=" * 60)
print("【樂莫集團分潤方式.xlsx】")
print("=" * 60)
try:
    df = pd.read_excel(r'C:\Users\User\Desktop\樂莫集團分潤方式.xlsx', sheet_name=None)
    for name, sheet in df.items():
        print(f"\n--- Sheet: {name} ---")
        print(sheet.to_string())
except Exception as e:
    print(f"Error: {e}")

print("\n" + "=" * 60)
print("【樂莫集團租稅規劃.docx】")
print("=" * 60)
try:
    doc = Document(r'C:\Users\User\Desktop\樂莫集團稅務制度\樂莫集團租稅規劃.docx')
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
except Exception as e:
    print(f"Error: {e}")

print("\n" + "=" * 60)
print("【樂莫集團架構.docx】")
print("=" * 60)
try:
    doc = Document(r'C:\Users\User\Desktop\樂莫集團架構\樂莫集團架構.docx')
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
except Exception as e:
    print(f"Error: {e}")
